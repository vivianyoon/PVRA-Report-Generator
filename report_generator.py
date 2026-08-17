import io
import re
import zipfile
from typing import Dict, List

import pandas as pd
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

APP_TITLE = "PVRA Kobo Excel to Village Word Reports"
MAIN_SHEET_CANDIDATES = ["PVRA Finding", "PVRA Village Assessment Form"]
REPEAT_SHEETS = {
    "livelihood_repeat": "Livelihood Information",
    "agriculture_repeat": "Agri Information",
    "livestock_repeat": "Livestock Data",
    "hazard_repeat": "Hazard Information",
    "seasonal_calendar_repeat": "Seasonal Calendar Information",
    "priority_ranking_repeat": "Priority Ranking Information",
}

FIELD_ALIASES = {
    # Main row
    "village_name": ["village_name_mm", "village_name", "ကျေးရွာအမည်"],
    "district_name": ["district_name", "ခရိုင်အမည်"],
    "township_name": ["township_name", "မြို့နယ်အမည်"],
    "village_tract": ["village_tract", "ကျေးရွာအုပ်စု"],
    "assessment_date": ["assessment_date", "PVRA ရေးဆွဲသည့်နေ့"],
    "enumerator_name": ["enumerator_name", "ဒေတာကောက်ယူသူအမည်"],
    "supporting_organization": ["supporting_organization", "ပါဝင်ကူညီသည့်အဖွဲ့အစည်း(များ)"],
    "recorder_name": ["recorder_name", "မှတ်တမ်းတင်သူ"],
    "facilitator_name": ["facilitator_name", "ဆွေးနွေးပံ့ပိုးသူ/အဖွဲ့"],
    "village_name_meaning": ["village_name_meaning", "ကျေးရွာအမည်/အဓိပ္ပါယ်"],
    "location_north": ["location_north", "မြောက်ဘက်တည်နေရာ"],
    "location_south": ["location_south", "တောင်ဘက်တည်နေရာ"],
    "location_east": ["location_east", "အရှေ့ဘက်တည်နေရာ"],
    "location_west": ["location_west", "အနောက်ဘက်တည်နေရာ"],
    "water_stream_condition": ["water_stream_condition", "ချောင်းမြောင်းစီးဆင်းမှုအခြေအနေ"],
    "irrigation_available": ["irrigation_available", "ဆည်ရေရှိပါသလား?"],
    "irrigation_supported_acre": ["irrigation_supported_acre", "ဆည်ရေဖြင့်ပံ့ပိုးနိုင်သော စိုက်ဧက"],
    "road_bridge_description": ["road_bridge_description", "လမ်းတံတားအမျိုးအစား၊ အရေအတွက်၊ အခြေအနေ"],
    "governance_description": ["governance_description", "အုပ်ချုပ်ရေးအခြေအနေ"],
    "accountability_description": ["accountability_description", "လူထုအားစောင့်ရှောက်မှု၊ တာဝန်ခံမှု"],
    "climate_description": ["climate_description", "ရာသီဥတုအခြေအနေ"],
    "avg_temperature": ["avg_temperature", "ပျမ်းမျှအပူချိန် (°C)"],
    "rainfall_description": ["rainfall_description", "မိုးရွာသွန်းမှုအခြေအနေ"],
    "forest_type_area": ["forest_type_area", "သစ်တော/စိုက်ခင်းအမျိုးအစားနှင့် အကျယ်အဝန်း"],
    "natural_resources": ["natural_resources", "သဘာဝသယံဇာတအမျိုးအစား၊ အခြေအနေ၊ ထုတ်ယူမှု"],
    "biodiversity": ["biodiversity", "ဇီဝမျိုးကွဲအမျိုးများ၊ အခြေအနေ၊ ထုတ်ယူမှု"],
    "total_population": ["total_population", "လူဦးရေစုစုပေါင်း"],
    "male_population": ["male_population", "အမျိုးသားဦးရေ"],
    "female_population": ["female_population", "အမျိုးသမီးဦးရေ"],
    "household_count": ["household_count", "အိမ်ထောင်စုအရေအတွက်"],
    "house_count": ["house_count", "အိမ်အရေအတွက်"],
    "education_facility_desc": ["education_facility_desc", "ပညာသင်ကျောင်းအခြေအနေ"],
    "school_count": ["school_count", "ကျောင်းအရေအတွက်"],
    "teacher_count": ["teacher_count", "ဆရာ/မ အရေအတွက်"],
    "student_count": ["student_count", "ကျောင်းသား/သူ အရေအတွက်"],
    "health_facility_desc": ["health_facility_desc", "ကျန်းမာရေးဌာနအမျိုးအစားနှင့် ရရှိသည့်ဝန်ဆောင်မှု"],
    "health_staff_count": ["health_staff_count", "ကျန်းမာရေးဝန်ထမ်းအရေအတွက်"],
    "education_committee_exists": ["education_committee_exists", "ပညာရေးအဖွဲ့ ရှိပါသလား?"],
    "health_committee_exists": ["health_committee_exists", "ကျန်းမာရေးအဖွဲ့ ရှိပါသလား?"],
    "social_orgs_description": ["social_orgs_description", "လူမှုရေးအသင်းအဖွဲ့များအခြေအနေ၊ လည်ပတ်ပုံ၊ ကောင်းကျိုးများ"],
    "electricity_status": ["electricity_status", "လျှပ်စစ်မီးအခြေအနေ"],
    "drinking_water_status": ["drinking_water_status", "သောက်ရေအခြေအနေ"],
    "sanitation_status": ["sanitation_status", "မိလ္လာအခြေအနေ"],
    "waste_management_status": ["waste_management_status", "အမှိုက်စီမံခန့်ခွဲမှုအခြေအနေ"],
    "public_services_description": ["public_services_description", "အခြားပြည်သူ့ရေးရာဝန်ဆောင်မှုအခြေအနေ"],
    "livelihood_summary": ["livelihood_summary", "အသက်မွေးလုပ်ငန်းအကျဉ်းချုပ်"],
    "agriculture_summary": ["agriculture_summary", "စိုက်ပျိုးရေးအကျဉ်းချုပ်"],
    "livestock_summary": ["livestock_summary", "မွေးမြူရေးအကျဉ်းချုပ်"],
    "general_remarks": ["general_remarks", "အထွေထွေမှတ်ချက်"],

    # Livelihood repeat
    "livelihood_type": ["livelihood_type", "အသက်မွေးလုပ်ငန်းအမျိုးအစား"],
    "livelihood_name": ["livelihood_name", "လုပ်ငန်းအမည်"],
    "livelihood_household_involved": ["livelihood_household_involved", "ပါဝင်သော အိမ်ထောင်စုအရေအတွက်"],
    "estimated_income": ["estimated_income", "ခန့်မှန်းဝင်ငွေ"],
    "income_period": ["income_period", "ဝင်ငွေကာလ"],
    "livelihood_condition": ["livelihood_condition", "လုပ်ငန်းအခြေအနေ"],
    "livelihood_challenges": ["livelihood_challenges", "စိန်ခေါ်ချက်များ"],

    # Agriculture repeat
    "crop_type": ["crop_type", "သီးနှံအမျိုးအစား"],
    "land_area_acre": ["land_area_acre", "စိုက်ပျိုးမြေဧက"],
    "average_yield": ["average_yield", "ပျမ်းမျှထွက်နှုန်း"],
    "yield_unit": ["yield_unit", "ထွက်နှုန်းယူနစ်"],
    "season_type": ["season_type", "စိုက်ရာသီ"],
    "agriculture_challenges": ["agriculture_challenges", "စိုက်ပျိုးရေးစိန်ခေါ်ချက်များ"],

    # Livestock repeat
    "livestock_type": ["livestock_type", "မွေးမြူရေးအမျိုးအစား"],
    "average_output": ["average_output", "ထွက်နှုန်း/ရလဒ်"],
    "livestock_condition": ["livestock_condition", "မွေးမြူရေးအခြေအနေ"],
    "livestock_challenges": ["livestock_challenges", "မွေးမြူရေးစိန်ခေါ်ချက်များ"],

    # Hazard repeat and/or main row (new structure moves many fields to main)
    "hazard_type": ["hazard_type", "ဘေးဖြစ်ရပ်အမျိုးအစား"],
    "hazard_date_or_period": ["hazard_date_or_period", "ဖြစ်ပွားသည့်အချိန်"],
    "hazard_description": ["hazard_description", "ဘေးဖြစ်ရပ်အကြောင်းအရာ"],
    "physical_impact": ["physical_impact", "ရုပ်ပိုင်းထိခိုက်မှု"],
    "psychosocial_impact": ["psychosocial_impact", "စိတ်ပိုင်းထိခိုက်မှု"],
    "economic_impact": ["economic_impact", "အလုပ်အကိုင်နှင့် စီးပွားရေးထိခိုက်မှု"],
    "environmental_impact": ["environmental_impact", "သဘာဝပတ်ဝန်းကျင်ထိခိုက်မှု"],
    "land_water_impact": ["land_water_impact", "မြေနှင့် ရေအရင်းအမြစ်ထိခိုက်မှု"],
    "household_expenditure_impact": ["household_expenditure_impact", "မိသားစုအသုံးစရိတ်အပေါ်ထိခိုက်မှု"],
    "food_availability_status": ["food_availability_status", "၁၂ လအတွင်း စားနပ်ရိက္ခာရရှိမှုအခြေအနေ"],
    "food_insecure_months": ["food_insecure_months", "စားနပ်ရိက္ခာမလုံလောက်သည့်လများ"],
    "coping_mechanisms": ["coping_mechanisms", "ဖြေရှင်းနည်း/ကျော်လွှားနည်း"],
    "agriculture_impact": ["agriculture_impact", "စိုက်ပျိုးရေးလုပ်ငန်းများအပေါ်ထိခိုက်မှု"],
    "agri_yield_status": ["agri_yield_status", "၁၂ လအတွင်း စိုက်ပျိုးထွက်နှုန်းအခြေအနေ", "စိုက်ပျိုးထွက်နှုန်းအခြေအနေ"],
    "agri_yield_reasons": ["agri_yield_reasons", "စိုက်ပျိုးထွက်နှုန်းနည်းရသည့်အကြောင်းရင်းများ"],
    "agri_vulnerability_factors": ["agri_vulnerability_factors", "စိုက်ပျိုးရေးအပေါ်ထိခိုက်လွယ်သည့်အဓိကအချက်များ"],
    "agri_most_vulnerable_groups": ["agri_most_vulnerable_groups", "ထိခိုက်လွယ်သည့်တောင်သူအမျိုးအစားများ"],
    "livestock_impact": ["livestock_impact", "မွေးမြူရေးလုပ်ငန်းများအပေါ်ထိခိုက်မှု"],
    "livestock_yield_status": ["livestock_yield_status", "၁၂ လအတွင်း မွေးမြူရေးထွက်နှုန်းအခြေအနေ", "မွေးမြူရေးထွက်နှုန်းအခြေအနေ"],
    "livestock_yield_reasons": ["livestock_yield_reasons", "မွေးမြူရေးထွက်နှုန်းနည်းရသည့်အကြောင်းရင်းများ"],
    "livestock_vulnerability_factors": ["livestock_vulnerability_factors", "မွေးမြူရေးအပေါ်ထိခိုက်လွယ်သည့်အဓိကအချက်များ"],
    "vulnerable_group_impact": ["vulnerable_group_impact", "အမျိုးသမီး/ကလေး/မသန်စွမ်း/သက်ကြီးရွယ်အိုများအပေါ်ထိခိုက်မှု"],
    "vulnerable_household_count": ["vulnerable_household_count", "ထိခိုက်လွယ်ဆုံးသောအိမ်ထောင်စုအရေအတွက်"],
    "vulnerable_root_causes": ["vulnerable_root_causes", "ထိခိုက်နိုင်သည့်အခြေခံအကြောင်းရင်းများ"],
    "community_group_impact": ["community_group_impact", "ရပ်ရွာလူမှုအဖွဲ့များအပေါ်ထိခိုက်မှု"],
    "community_recovery_actions": ["community_recovery_actions", "ဘယ်လိုထူထောင်/ကျော်လွှားခဲ့ကြသလဲ"],
    "climate_change_pattern": ["climate_change_pattern", "ရာသီဥတုပြောင်းလဲမှုပုံစံ"],
    "pest_disease_pattern": ["pest_disease_pattern", "ပိုးမွှား/ရောဂါပုံစံနှင့် အကြောင်းရင်း"],

    # Seasonal calendar
    "calendar_category": ["calendar_category", "အမျိုးအစား"],
    "calendar_item_name": ["calendar_item_name", "အကြောင်းအရာအမည်"],
    "active_months": ["active_months", "သက်ဆိုင်သည့်လများ"],
    "peak_months": ["peak_months", "အများဆုံး/အမြင့်ဆုံးဖြစ်သည့်လများ"],
    "low_months": ["low_months", "အနည်းဆုံး/အနိမ့်ဆုံးဖြစ်သည့်လများ"],
    "calendar_notes": ["calendar_notes", "ရှင်းလင်းချက်"],

    # Ranking
    "ranking_category": ["ranking_category", "ကဏ္ဍ"],
    "option_name": ["option_name", "လုပ်ငန်း/အဆိုပြုချက်အမည်"],
    "score": ["score", "ရမှတ်"],
    "rank_order": ["rank_order", "အဆင့်"],
    "selected_top3": ["selected_top3", "Top 3 ထဲပါသလား?"],
    "needs_relevance": ["needs_relevance", "လိုအပ်ချက်နှင့်ကိုက်ညီမှု"],
    "feasibility": ["feasibility", "ဖြစ်နိုင်ချေ"],
    "practicality": ["practicality", "လက်တွေ့ကျမှု"],
    "vulnerable_inclusion": ["vulnerable_inclusion", "ထိခိုက်လွယ်သူများထည့်သွင်းစဉ်းစားမှု"],
    "short_term_feasible": ["short_term_feasible", "အချိန်တိုကာလအတွင်း အကောင်အထည်ဖော်နိုင်ပါသလား?"],
    "social_protection_link": ["social_protection_link", "လူမှုကာကွယ်စောင့်ရှောက်ခြင်းဆိုင်ရာချိတ်ဆက်မှု"],
    "community_feedback": ["community_feedback", "လူထုသဘောထား"],
}

EN_MONTH_MAP = {"jan": "Jan", "feb": "Feb", "mar": "Mar", "apr": "Apr", "may": "May", "jun": "Jun", "jul": "Jul", "aug": "Aug", "sep": "Sep", "oct": "Oct", "nov": "Nov", "dec": "Dec"}
MY_MONTH_MAP = {
    "ဇန်နဝါရီ": "ဇန်နဝါရီ", "ဖေဖော်ဝါရီ": "ဖေဖော်ဝါရီ", "မတ်": "မတ်", "ဧပြီ": "ဧပြီ", "မေ": "မေ", "ဇွန်": "ဇွန်",
    "ဇူလိုင်": "ဇူလိုင်", "သြဂုတ်": "သြဂုတ်", "စက်တင်ဘာ": "စက်တင်ဘာ", "အောက်တိုဘာ": "အောက်တိုဘာ", "နိုဝင်ဘာ": "နိုဝင်ဘာ", "ဒီဇင်ဘာ": "ဒီဇင်ဘာ"
}
MONTH_NAME_MAP = {**EN_MONTH_MAP, **MY_MONTH_MAP}

VALUE_LABELS = {
    "yes": "Yes", "no": "No", "none": "None", "poor": "Poor", "fair": "Fair", "good": "Good",
    "normal": "Normal", "decreased": "Decreased", "increased": "Increased", "seasonal": "Seasonal",
    "monthly": "Monthly", "daily": "Daily", "weekly": "Weekly", "trade": "Trade", "casual_labor": "Casual labor",
    "crop_activity": "Crop activity", "livelihood": "Livelihood", "livestock": "Livestock", "agriculture": "Agriculture",
    "medium": "Medium", "high": "High", "low": "Low", "partial": "Partial", "pest": "Pest", "monsoon": "Monsoon",
    "winter": "Winter", "summer": "Summer",
    "ရှိ/ဟုတ်": "ရှိ/ဟုတ်", "မရှိ/မဟုတ်": "မရှိ/မဟုတ်", "မရှိ": "မရှိ", "ညံ့": "ညံ့", "အသင့်အတင့်": "အသင့်အတင့်",
    "ကောင်း": "ကောင်း", "ပုံမှန်": "ပုံမှန်", "လျော့": "လျော့", "တိုး": "တိုး",
}

TEXT = {
    "en": {
        "report_title": "Participatory Village Risk Assessment (PVRA) Report",
        "subtitle": "Village: {village} | Township: {township}",
        "no_data": "No data available.",
        "sec1": "1. Village General Information",
        "sec2": "2. Livelihood Information",
        "sec3": "3. Agri Information",
        "sec4": "4. Livestock Data",
        "sec5": "5. Hazard Information",
        "sec6": "6. Seasonal Calendar Information",
        "sec7": "7. Priority Ranking Information",
        "location": "Location and Natural Setting",
        "population": "Population and Social Services",
        "infrastructure": "Infrastructure and Public Services",
        "hazard_summary": "Village-level hazard and vulnerability summary",
        "livelihood_activities": "Livelihood activities",
        "livelihood_conditions": "Livelihood conditions and challenges",
        "agri_activities": "Agriculture activities",
        "agri_challenges": "Agriculture challenges",
        "livestock_activities": "Livestock activities",
        "livestock_conditions": "Livestock conditions and challenges",
        "hazard": "Hazard {i}: {name}",
        "item": "Item {i}",
        "crop": "Crop {i}",
        "category": "Category",
        "item_name": "Item name",
        "active_months": "Active months",
        "peak_months": "Peak months",
        "low_months": "Low months",
        "notes": "Notes",
        "option": "Option",
        "score": "Score",
        "rank": "Rank",
        "top3": "Top 3",
        "needs_relevance": "Needs relevance",
        "feasibility": "Feasibility",
        "practicality": "Practicality",
        "vulnerable_inclusion": "Vulnerable inclusion",
        "short_term_feasible": "Short-term feasible",
        "social_protection_link": "Social protection link",
        "community_feedback": "Community feedback",
        "field": {
            "Village Name": "Village Name", "Village Tract": "Village Tract", "Township": "Township", "District": "District",
            "Assessment Date": "Assessment Date", "Enumerator": "Enumerator", "Facilitator": "Facilitator", "Recorder": "Recorder",
            "Supporting Organization": "Supporting Organization", "Village name meaning": "Village name meaning", "North": "North",
            "South": "South", "East": "East", "West": "West", "Water stream condition": "Water stream condition",
            "Irrigation available": "Irrigation available", "Irrigation supported acre": "Irrigation supported acre",
            "Climate description": "Climate description", "Average temperature": "Average temperature", "Rainfall description": "Rainfall description",
            "Forest type / area": "Forest type / area", "Natural resources": "Natural resources", "Biodiversity": "Biodiversity",
            "Total population": "Total population", "Male population": "Male population", "Female population": "Female population",
            "Household count": "Household count", "House count": "House count", "Education facility": "Education facility",
            "School count": "School count", "Teacher count": "Teacher count", "Student count": "Student count",
            "Health facility": "Health facility", "Health staff count": "Health staff count", "Education committee exists": "Education committee exists",
            "Health committee exists": "Health committee exists", "Social organizations": "Social organizations", "Road and bridge description": "Road and bridge description",
            "Governance description": "Governance description", "Accountability description": "Accountability description", "Electricity status": "Electricity status",
            "Drinking water status": "Drinking water status", "Sanitation status": "Sanitation status", "Waste management status": "Waste management status",
            "Public services description": "Public services description", "General remarks": "General remarks", "Livelihood summary": "Livelihood summary",
            "Agriculture summary": "Agriculture summary", "Livestock summary": "Livestock summary", "Type": "Type", "Livelihood name": "Livelihood name",
            "HH involved": "HH involved", "Estimated income": "Estimated income", "Income period": "Income period", "Condition": "Condition",
            "Challenges": "Challenges", "Crop type": "Crop type", "Land area (acre)": "Land area (acre)", "Average yield": "Average yield",
            "Yield unit": "Yield unit", "Season type": "Season type", "Livestock type": "Livestock type", "Average output": "Average output",
            "Hazard type": "Hazard type", "Date / period": "Date / period", "Food availability status": "Food availability status",
            "Food insecure months": "Food insecure months", "Agri yield status": "Agri yield status", "Agri yield reasons": "Reasons for low agri yield",
            "Livestock yield status": "Livestock yield status", "Livestock yield reasons": "Reasons for low livestock yield", "Vulnerable household count": "Vulnerable household count", "Hazard description": "Hazard description", "Physical impact": "Physical impact",
            "Psychosocial impact": "Psychosocial impact", "Economic impact": "Economic impact", "Environmental impact": "Environmental impact",
            "Land / water impact": "Land / water impact", "Household expenditure impact": "Household expenditure impact", "Coping mechanisms": "Coping mechanisms",
            "Agriculture impact": "Agriculture impact", "Agriculture vulnerability factors": "Agriculture vulnerability factors", "Most vulnerable groups in agriculture": "Most vulnerable groups in agriculture",
            "Livestock impact": "Livestock impact", "Livestock vulnerability factors": "Livestock vulnerability factors", "Vulnerable group impact": "Vulnerable group impact",
            "Vulnerable root causes": "Vulnerable root causes", "Community group impact": "Community group impact", "Community recovery actions": "Community recovery actions",
            "Climate change pattern": "Climate change pattern", "Pest / disease pattern": "Pest / disease pattern"
        }
    },
    "my": {
        "report_title": "ကျေးရွာအခြေပြု ထိခိုက်လွယ်မှုနှင့် အန္တရာယ်ဆန်းစစ်ခြင်း (PVRA) အစီရင်ခံစာ",
        "subtitle": "ကျေးရွာ - {village} | မြို့နယ် - {township}",
        "no_data": "ဒေတာမရှိပါ။",
        "sec1": "၁။ ကျေးရွာအထွေထွေအချက်အလက်",
        "sec2": "၂။ အသက်မွေးလုပ်ငန်းဆိုင်ရာအချက်အလက်",
        "sec3": "၃။ စိုက်ပျိုးရေးဆိုင်ရာအချက်အလက်",
        "sec4": "၄။ မွေးမြူရေးဆိုင်ရာအချက်အလက်",
        "sec5": "၅။ ဘေးအန္တရာယ်ဆိုင်ရာအချက်အလက်",
        "sec6": "၆။ ရာသီပြက္ခဒိန်ဆိုင်ရာအချက်အလက်",
        "sec7": "၇။ ဦးစားပေးအဆင့်သတ်မှတ်ချက်ဆိုင်ရာအချက်အလက်",
        "location": "တည်နေရာနှင့် သဘာဝပတ်ဝန်းကျင်",
        "population": "လူဦးရေ၊ လူမှုရေးနှင့် အခြေခံဝန်ဆောင်မှုများ",
        "infrastructure": "အခြေခံအဆောက်အဦနှင့် အများပြည်သူဝန်ဆောင်မှုများ",
        "hazard_summary": "ကျေးရွာအဆင့် ဘေးအန္တရာယ်နှင့် ထိခိုက်လွယ်မှုအကျဉ်းချုပ်",
        "livelihood_activities": "အသက်မွေးလုပ်ငန်းအချက်အလက်များ",
        "livelihood_conditions": "အသက်မွေးလုပ်ငန်းအခြေအနေနှင့် စိန်ခေါ်ချက်များ",
        "agri_activities": "စိုက်ပျိုးရေးအချက်အလက်များ",
        "agri_challenges": "စိုက်ပျိုးရေးစိန်ခေါ်ချက်များ",
        "livestock_activities": "မွေးမြူရေးအချက်အလက်များ",
        "livestock_conditions": "မွေးမြူရေးအခြေအနေနှင့် စိန်ခေါ်ချက်များ",
        "hazard": "ဘေးဖြစ်ရပ် {i}: {name}",
        "item": "အကြောင်းအရာ {i}",
        "crop": "သီးနှံ {i}",
        "category": "အမျိုးအစား",
        "item_name": "အကြောင်းအရာအမည်",
        "active_months": "သက်ဆိုင်သည့်လများ",
        "peak_months": "အများဆုံး/အမြင့်ဆုံးဖြစ်သည့်လများ",
        "low_months": "အနည်းဆုံး/အနိမ့်ဆုံးဖြစ်သည့်လများ",
        "notes": "ရှင်းလင်းချက်",
        "option": "လုပ်ငန်း/အဆိုပြုချက်အမည်",
        "score": "ရမှတ်",
        "rank": "အဆင့်",
        "top3": "Top 3 ထဲပါသလား?",
        "needs_relevance": "လိုအပ်ချက်နှင့်ကိုက်ညီမှု",
        "feasibility": "ဖြစ်နိုင်ချေ",
        "practicality": "လက်တွေ့ကျမှု",
        "vulnerable_inclusion": "ထိခိုက်လွယ်သူများထည့်သွင်းစဉ်းစားမှု",
        "short_term_feasible": "အချိန်တိုကာလအတွင်း အကောင်အထည်ဖော်နိုင်မှု",
        "social_protection_link": "လူမှုကာကွယ်စောင့်ရှောက်ခြင်းဆိုင်ရာချိတ်ဆက်မှု",
        "community_feedback": "လူထုသဘောထား",
        "field": {
            "Village Name": "ကျေးရွာအမည်", "Village Tract": "ကျေးရွာအုပ်စု", "Township": "မြို့နယ်အမည်", "District": "ခရိုင်အမည်",
            "Assessment Date": "PVRA ရေးဆွဲသည့်နေ့", "Enumerator": "ဒေတာကောက်ယူသူအမည်", "Facilitator": "ဆွေးနွေးပံ့ပိုးသူ/အဖွဲ့", "Recorder": "မှတ်တမ်းတင်သူ",
            "Supporting Organization": "ပါဝင်ကူညီသည့်အဖွဲ့အစည်း(များ)", "Village name meaning": "ကျေးရွာအမည်/အဓိပ္ပါယ်", "North": "မြောက်ဘက်တည်နေရာ",
            "South": "တောင်ဘက်တည်နေရာ", "East": "အရှေ့ဘက်တည်နေရာ", "West": "အနောက်ဘက်တည်နေရာ", "Water stream condition": "ချောင်းမြောင်းစီးဆင်းမှုအခြေအနေ",
            "Irrigation available": "ဆည်ရေရှိပါသလား?", "Irrigation supported acre": "ဆည်ရေဖြင့်ပံ့ပိုးနိုင်သော စိုက်ဧက", "Climate description": "ရာသီဥတုအခြေအနေ",
            "Average temperature": "ပျမ်းမျှအပူချိန် (°C)", "Rainfall description": "မိုးရွာသွန်းမှုအခြေအနေ", "Forest type / area": "သစ်တော/စိုက်ခင်းအမျိုးအစားနှင့် အကျယ်အဝန်း",
            "Natural resources": "သဘာဝသယံဇာတအမျိုးအစား၊ အခြေအနေ၊ ထုတ်ယူမှု", "Biodiversity": "ဇီဝမျိုးကွဲအမျိုးများ၊ အခြေအနေ၊ ထုတ်ယူမှု", "Total population": "လူဦးရေစုစုပေါင်း",
            "Male population": "အမျိုးသားဦးရေ", "Female population": "အမျိုးသမီးဦးရေ", "Household count": "အိမ်ထောင်စုအရေအတွက်", "House count": "အိမ်အရေအတွက်",
            "Education facility": "ပညာသင်ကျောင်းအခြေအနေ", "School count": "ကျောင်းအရေအတွက်", "Teacher count": "ဆရာ/မ အရေအတွက်", "Student count": "ကျောင်းသား/သူ အရေအတွက်",
            "Health facility": "ကျန်းမာရေးဌာနအမျိုးအစားနှင့် ရရှိသည့်ဝန်ဆောင်မှု", "Health staff count": "ကျန်းမာရေးဝန်ထမ်းအရေအတွက်", "Education committee exists": "ပညာရေးအဖွဲ့ ရှိပါသလား?",
            "Health committee exists": "ကျန်းမာရေးအဖွဲ့ ရှိပါသလား?", "Social organizations": "လူမှုရေးအသင်းအဖွဲ့များအခြေအနေ၊ လည်ပတ်ပုံ၊ ကောင်းကျိုးများ", "Road and bridge description": "လမ်းတံတားအမျိုးအစား၊ အရေအတွက်၊ အခြေအနေ",
            "Governance description": "အုပ်ချုပ်ရေးအခြေအနေ", "Accountability description": "လူထုအားစောင့်ရှောက်မှု၊ တာဝန်ခံမှု", "Electricity status": "လျှပ်စစ်မီးအခြေအနေ", "Drinking water status": "သောက်ရေအခြေအနေ",
            "Sanitation status": "မိလ္လာအခြေအနေ", "Waste management status": "အမှိုက်စီမံခန့်ခွဲမှုအခြေအနေ", "Public services description": "အခြားပြည်သူ့ရေးရာဝန်ဆောင်မှုအခြေအနေ",
            "General remarks": "အထွေထွေမှတ်ချက်", "Livelihood summary": "အသက်မွေးလုပ်ငန်းအကျဉ်းချုပ်", "Agriculture summary": "စိုက်ပျိုးရေးအကျဉ်းချုပ်", "Livestock summary": "မွေးမြူရေးအကျဉ်းချုပ်",
            "Type": "အသက်မွေးလုပ်ငန်းအမျိုးအစား", "Livelihood name": "လုပ်ငန်းအမည်", "HH involved": "ပါဝင်သော အိမ်ထောင်စုအရေအတွက်", "Estimated income": "ခန့်မှန်းဝင်ငွေ",
            "Income period": "ဝင်ငွေကာလ", "Condition": "အခြေအနေ", "Challenges": "စိန်ခေါ်ချက်များ", "Crop type": "သီးနှံအမျိုးအစား", "Land area (acre)": "စိုက်ပျိုးမြေဧက",
            "Average yield": "ပျမ်းမျှထွက်နှုန်း", "Yield unit": "ထွက်နှုန်းယူနစ်", "Season type": "စိုက်ရာသီ", "Livestock type": "မွေးမြူရေးအမျိုးအစား", "Average output": "ထွက်နှုန်း/ရလဒ်",
            "Hazard type": "ဘေးဖြစ်ရပ်အမျိုးအစား", "Date / period": "ဖြစ်ပွားသည့်အချိန်", "Food availability status": "၁၂ လအတွင်း စားနပ်ရိက္ခာရရှိမှုအခြေအနေ",
            "Food insecure months": "စားနပ်ရိက္ခာမလုံလောက်သည့်လများ", "Agri yield status": "၁၂ လအတွင်း စိုက်ပျိုးထွက်နှုန်းအခြေအနေ", "Agri yield reasons": "စိုက်ပျိုးထွက်နှုန်းနည်းရသည့်အကြောင်းရင်းများ",
            "Livestock yield status": "၁၂ လအတွင်း မွေးမြူရေးထွက်နှုန်းအခြေအနေ", "Livestock yield reasons": "မွေးမြူရေးထွက်နှုန်းနည်းရသည့်အကြောင်းရင်းများ", "Vulnerable household count": "ထိခိုက်လွယ်ဆုံးသောအိမ်ထောင်စုအရေအတွက်", "Hazard description": "ဘေးဖြစ်ရပ်အကြောင်းအရာ", "Physical impact": "ရုပ်ပိုင်းထိခိုက်မှု",
            "Psychosocial impact": "စိတ်ပိုင်းထိခိုက်မှု", "Economic impact": "အလုပ်အကိုင်နှင့် စီးပွားရေးထိခိုက်မှု", "Environmental impact": "သဘာဝပတ်ဝန်းကျင်ထိခိုက်မှု", "Land / water impact": "မြေနှင့် ရေအရင်းအမြစ်ထိခိုက်မှု",
            "Household expenditure impact": "မိသားစုအသုံးစရိတ်အပေါ်ထိခိုက်မှု", "Coping mechanisms": "ဖြေရှင်းနည်း/ကျော်လွှားနည်း", "Agriculture impact": "စိုက်ပျိုးရေးလုပ်ငန်းများအပေါ်ထိခိုက်မှု",
            "Agriculture vulnerability factors": "စိုက်ပျိုးရေးအပေါ်ထိခိုက်လွယ်သည့်အဓိကအချက်များ", "Most vulnerable groups in agriculture": "ထိခိုက်လွယ်သည့်တောင်သူအမျိုးအစားများ", "Livestock impact": "မွေးမြူရေးလုပ်ငန်းများအပေါ်ထိခိုက်မှု",
            "Livestock vulnerability factors": "မွေးမြူရေးအပေါ်ထိခိုက်လွယ်သည့်အဓိကအချက်များ", "Vulnerable group impact": "အမျိုးသမီး/ကလေး/မသန်စွမ်း/သက်ကြီးရွယ်အိုများအပေါ်ထိခိုက်မှု", "Vulnerable root causes": "ထိခိုက်နိုင်သည့်အခြေခံအကြောင်းရင်းများ",
            "Community group impact": "ရပ်ရွာလူမှုအဖွဲ့များအပေါ်ထိခိုက်မှု", "Community recovery actions": "ဘယ်လိုထူထောင်/ကျော်လွှားခဲ့ကြသလဲ", "Climate change pattern": "ရာသီဥတုပြောင်းလဲမှုပုံစံ", "Pest / disease pattern": "ပိုးမွှား/ရောဂါပုံစံနှင့် အကြောင်းရင်း"
        }
    }
}


def detect_language(df: pd.DataFrame) -> str:
    cols = set(map(str, df.columns))
    return "my" if "ကျေးရွာအမည်" in cols else "en"


def tr(lang: str, key: str) -> str:
    return TEXT[lang].get(key, key)


def label(lang: str, key: str) -> str:
    return TEXT[lang]["field"].get(key, key)


def getv(obj, canonical: str, default="-"):
    aliases = FIELD_ALIASES.get(canonical, [canonical])
    if hasattr(obj, "index"):
        for col in aliases:
            if col in obj.index and not is_blank(obj[col]):
                return obj[col]
    return default


def getcol(df: pd.DataFrame, canonical: str):
    for col in FIELD_ALIASES.get(canonical, [canonical]):
        if col in df.columns:
            return col
    return None


def find_main_sheet(workbook: Dict[str, pd.DataFrame]) -> str | None:
    for name in MAIN_SHEET_CANDIDATES:
        if name in workbook:
            return name
    return None


def sanitize_filename(value: str) -> str:
    value = str(value or "village_report").strip()
    value = re.sub(r"[\\/:*?\"<>|]", "_", value)
    value = re.sub(r"\s+", " ", value).strip()
    return value[:120] or "village_report"


def is_blank(value) -> bool:
    if pd.isna(value):
        return True
    if isinstance(value, str) and value.strip() == "":
        return True
    return False


def fmt_value(value):
    if pd.isna(value):
        return "-"
    if isinstance(value, pd.Timestamp):
        return value.strftime("%d-%b-%Y")
    if isinstance(value, float) and value.is_integer():
        return f"{int(value):,}"
    if isinstance(value, (int, float)):
        return f"{value:,}"
    if isinstance(value, str):
        clean = value.strip()
        return VALUE_LABELS.get(clean, clean)
    return str(value)


def clean_text(value):
    text = fmt_value(value)
    return "-" if text == "-" else str(text).replace("_", " ")


def month_list_from_row(row, canonical_prefix: str):
    base_cols = FIELD_ALIASES.get(canonical_prefix, [canonical_prefix])
    selected = []
    seen = set()
    for col in row.index:
        col_str = str(col)
        for base in base_cols:
            if col_str.startswith(f"{base}/") and not is_blank(row[col]):
                try:
                    val = int(row[col])
                except Exception:
                    val = 0
                if val == 1:
                    suffix = col_str.split("/", 1)[1].strip()
                    month_label = MONTH_NAME_MAP.get(suffix.lower(), MONTH_NAME_MAP.get(suffix, suffix))
                    if month_label not in seen:
                        selected.append(month_label)
                        seen.add(month_label)
    if not selected:
        raw = getv(row, canonical_prefix, default="")
        if raw not in ["", "-"]:
            for token in str(raw).split():
                token = token.strip()
                if token:
                    selected.append(MONTH_NAME_MAP.get(token.lower(), MONTH_NAME_MAP.get(token, token)))
    return ", ".join(selected) if selected else "-"


def remove_metadata(df: pd.DataFrame) -> pd.DataFrame:
    return df[[c for c in df.columns if not str(c).startswith("_") and not str(c).startswith("meta/") and not str(c).endswith("_URL")]].copy()


def set_cell_text(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(str(text))
    run.bold = bold
    run.font.name = "Noto Sans Myanmar"
    run.font.size = Pt(9.5)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for side in ["top", "start", "bottom", "end"]:
        node = tcMar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tcMar.append(node)
        node.set(qn("w:w"), "90")
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)


def style_table(table, header_fill="D9EAF7"):
    table.style = "Table Grid"
    table.autofit = True
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for cell in hdr.cells:
        shade_cell(cell, header_fill)
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.bold = True
                run.font.name = "Noto Sans Myanmar"
                run.font.size = Pt(9.5)
    for row in table.rows[1:]:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = "Noto Sans Myanmar"
                    run.font.size = Pt(9.5)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    if level == 0:
        p.style = doc.styles["Title"]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.name = "Noto Sans Myanmar"
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = RGBColor(31, 78, 121)
    else:
        p.style = doc.styles["Heading 1"] if level == 1 else doc.styles["Heading 2"]
        run = p.add_run(text)
        run.font.name = "Noto Sans Myanmar"
        run.font.bold = True
        run.font.color.rgb = RGBColor(31, 78, 121) if level == 1 else RGBColor(68, 68, 68)
    return p


def add_kv_table(doc, pairs, widths=(2.35, 4.65)):
    table = doc.add_table(rows=0, cols=2)
    table.autofit = False
    for key, value in pairs:
        row = table.add_row().cells
        row[0].width = Inches(widths[0])
        row[1].width = Inches(widths[1])
        set_cell_text(row[0], key, bold=True)
        shade_cell(row[0], "EAF2F8")
        set_cell_text(row[1], clean_text(value))
    style_table(table)
    doc.add_paragraph()


def add_df_table(doc, df: pd.DataFrame, title: str | None = None, empty_text="No data available."):
    if df.empty:
        doc.add_paragraph(empty_text)
        return
    if title:
        add_heading(doc, title, level=2)
    display_df = df.copy().fillna("-")
    table = doc.add_table(rows=1, cols=len(display_df.columns))
    for i, col in enumerate(display_df.columns):
        set_cell_text(table.rows[0].cells[i], str(col), bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    for _, row in display_df.iterrows():
        cells = table.add_row().cells
        for i, value in enumerate(row.tolist()):
            align = WD_ALIGN_PARAGRAPH.CENTER if isinstance(value, (int, float)) and not isinstance(value, bool) else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(cells[i], clean_text(value), align=align)
    style_table(table)
    doc.add_paragraph()


def add_paragraph_text(doc, field_label: str, value):
    p = doc.add_paragraph()
    r1 = p.add_run(f"{field_label}: ")
    r1.bold = True
    r1.font.name = "Noto Sans Myanmar"
    r1.font.size = Pt(10)
    r2 = p.add_run(clean_text(value))
    r2.font.name = "Noto Sans Myanmar"
    r2.font.size = Pt(10)


def base_document():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.7)
    sec.bottom_margin = Inches(0.65)
    sec.left_margin = Inches(0.75)
    sec.right_margin = Inches(0.75)
    for style_name in ["Normal", "Title", "Heading 1", "Heading 2"]:
        style = doc.styles[style_name]
        style.font.name = "Noto Sans Myanmar"
    doc.styles["Normal"].font.size = Pt(10)
    doc.styles["Heading 1"].font.size = Pt(13)
    doc.styles["Heading 2"].font.size = Pt(11)
    return doc


def load_workbook(uploaded_file) -> Dict[str, pd.DataFrame]:
    workbook = pd.read_excel(uploaded_file, sheet_name=None)
    for _, df in workbook.items():
        if "_parent_index" in df.columns:
            df["_parent_index"] = pd.to_numeric(df["_parent_index"], errors="coerce")
        if "_index" in df.columns:
            df["_index"] = pd.to_numeric(df["_index"], errors="coerce")
    return workbook


def village_name(row):
    return clean_text(getv(row, "village_name", default="Village"))


def local_df(df: pd.DataFrame, column_map: List[tuple], lang: str) -> pd.DataFrame:
    selected = {}
    for canonical, display_key in column_map:
        col = getcol(df, canonical)
        if col:
            selected[label(lang, display_key)] = df[col]
    return remove_metadata(pd.DataFrame(selected)) if selected else pd.DataFrame()


def section_village_general(doc, row, lang: str):
    add_heading(doc, tr(lang, "sec1"), level=1)
    add_kv_table(doc, [
        (label(lang, "Village Name"), getv(row, "village_name")),
        (label(lang, "Village Tract"), getv(row, "village_tract")),
        (label(lang, "Township"), getv(row, "township_name")),
        (label(lang, "District"), getv(row, "district_name")),
        (label(lang, "Assessment Date"), getv(row, "assessment_date")),
        (label(lang, "Enumerator"), getv(row, "enumerator_name")),
        (label(lang, "Facilitator"), getv(row, "facilitator_name")),
        (label(lang, "Recorder"), getv(row, "recorder_name")),
        (label(lang, "Supporting Organization"), getv(row, "supporting_organization")),
    ])
    add_heading(doc, tr(lang, "location"), level=2)
    add_kv_table(doc, [
        (label(lang, "Village name meaning"), getv(row, "village_name_meaning")),
        (label(lang, "North"), getv(row, "location_north")),
        (label(lang, "South"), getv(row, "location_south")),
        (label(lang, "East"), getv(row, "location_east")),
        (label(lang, "West"), getv(row, "location_west")),
        (label(lang, "Water stream condition"), getv(row, "water_stream_condition")),
        (label(lang, "Irrigation available"), getv(row, "irrigation_available")),
        (label(lang, "Irrigation supported acre"), getv(row, "irrigation_supported_acre")),
        (label(lang, "Climate description"), getv(row, "climate_description")),
        (label(lang, "Average temperature"), getv(row, "avg_temperature")),
        (label(lang, "Rainfall description"), getv(row, "rainfall_description")),
        (label(lang, "Forest type / area"), getv(row, "forest_type_area")),
        (label(lang, "Natural resources"), getv(row, "natural_resources")),
        (label(lang, "Biodiversity"), getv(row, "biodiversity")),
    ])
    add_heading(doc, tr(lang, "population"), level=2)
    add_kv_table(doc, [
        (label(lang, "Total population"), getv(row, "total_population")),
        (label(lang, "Male population"), getv(row, "male_population")),
        (label(lang, "Female population"), getv(row, "female_population")),
        (label(lang, "Household count"), getv(row, "household_count")),
        (label(lang, "House count"), getv(row, "house_count")),
        (label(lang, "Education facility"), getv(row, "education_facility_desc")),
        (label(lang, "School count"), getv(row, "school_count")),
        (label(lang, "Teacher count"), getv(row, "teacher_count")),
        (label(lang, "Student count"), getv(row, "student_count")),
        (label(lang, "Health facility"), getv(row, "health_facility_desc")),
        (label(lang, "Health staff count"), getv(row, "health_staff_count")),
        (label(lang, "Education committee exists"), getv(row, "education_committee_exists")),
        (label(lang, "Health committee exists"), getv(row, "health_committee_exists")),
        (label(lang, "Social organizations"), getv(row, "social_orgs_description")),
    ])
    add_heading(doc, tr(lang, "infrastructure"), level=2)
    add_kv_table(doc, [
        (label(lang, "Road and bridge description"), getv(row, "road_bridge_description")),
        (label(lang, "Governance description"), getv(row, "governance_description")),
        (label(lang, "Accountability description"), getv(row, "accountability_description")),
        (label(lang, "Electricity status"), getv(row, "electricity_status")),
        (label(lang, "Drinking water status"), getv(row, "drinking_water_status")),
        (label(lang, "Sanitation status"), getv(row, "sanitation_status")),
        (label(lang, "Waste management status"), getv(row, "waste_management_status")),
        (label(lang, "Public services description"), getv(row, "public_services_description")),
        (label(lang, "General remarks"), getv(row, "general_remarks")),
    ])


def section_livelihood(doc, row, child_df, lang: str):
    add_heading(doc, tr(lang, "sec2"), level=1)
    add_paragraph_text(doc, label(lang, "Livelihood summary"), getv(row, "livelihood_summary"))
    doc.add_paragraph()
    df = local_df(child_df, [
        ("livelihood_type", "Type"), ("livelihood_name", "Livelihood name"), ("livelihood_household_involved", "HH involved"),
        ("estimated_income", "Estimated income"), ("income_period", "Income period")
    ], lang)
    add_df_table(doc, df, title=tr(lang, "livelihood_activities"), empty_text=tr(lang, "no_data"))
    if not child_df.empty:
        add_heading(doc, tr(lang, "livelihood_conditions"), level=2)
        for i, (_, r) in enumerate(child_df.iterrows(), start=1):
            add_paragraph_text(doc, tr(lang, "item").format(i=i), f"{clean_text(getv(r, 'livelihood_name'))} ({clean_text(getv(r, 'livelihood_type'))})")
            add_paragraph_text(doc, label(lang, "Condition"), getv(r, "livelihood_condition"))
            add_paragraph_text(doc, label(lang, "Challenges"), getv(r, "livelihood_challenges"))
            doc.add_paragraph()


def section_agri(doc, row, child_df, lang: str):
    add_heading(doc, tr(lang, "sec3"), level=1)
    add_paragraph_text(doc, label(lang, "Agriculture summary"), getv(row, "agriculture_summary"))
    doc.add_paragraph()
    df = local_df(child_df, [
        ("crop_type", "Crop type"), ("land_area_acre", "Land area (acre)"), ("average_yield", "Average yield"),
        ("yield_unit", "Yield unit"), ("season_type", "Season type")
    ], lang)
    add_df_table(doc, df, title=tr(lang, "agri_activities"), empty_text=tr(lang, "no_data"))
    if not child_df.empty:
        add_heading(doc, tr(lang, "agri_challenges"), level=2)
        for i, (_, r) in enumerate(child_df.iterrows(), start=1):
            add_paragraph_text(doc, tr(lang, "crop").format(i=i), clean_text(getv(r, 'crop_type')))
            add_paragraph_text(doc, label(lang, "Challenges"), getv(r, "agriculture_challenges"))
            doc.add_paragraph()


def section_livestock(doc, row, child_df, lang: str):
    add_heading(doc, tr(lang, "sec4"), level=1)
    add_paragraph_text(doc, label(lang, "Livestock summary"), getv(row, "livestock_summary"))
    doc.add_paragraph()
    df = local_df(child_df, [
        ("livestock_type", "Livestock type"), ("livelihood_household_involved", "HH involved"), ("average_output", "Average output")
    ], lang)
    add_df_table(doc, df, title=tr(lang, "livestock_activities"), empty_text=tr(lang, "no_data"))
    if not child_df.empty:
        add_heading(doc, tr(lang, "livestock_conditions"), level=2)
        for i, (_, r) in enumerate(child_df.iterrows(), start=1):
            add_paragraph_text(doc, tr(lang, "item").format(i=i), clean_text(getv(r, 'livestock_type')))
            add_paragraph_text(doc, label(lang, "Condition"), getv(r, "livestock_condition"))
            add_paragraph_text(doc, label(lang, "Challenges"), getv(r, "livestock_challenges"))
            doc.add_paragraph()


def hazard_summary_value(main_row, child_df, canonical):
    main_value = getv(main_row, canonical, default="-")
    if main_value != "-":
        return main_value
    if child_df is not None and not child_df.empty:
        for _, child_row in child_df.iterrows():
            child_value = getv(child_row, canonical, default="-")
            if child_value != "-":
                return child_value
    return "-"


def section_hazard(doc, main_row, child_df, lang: str):
    add_heading(doc, tr(lang, "sec5"), level=1)

    summary_pairs = [
        (label(lang, "Food availability status"), hazard_summary_value(main_row, child_df, "food_availability_status")),
        (label(lang, "Food insecure months"), month_list_from_row(main_row if hazard_summary_value(main_row, child_df, "food_insecure_months") != "-" else (child_df.iloc[0] if not child_df.empty else main_row), "food_insecure_months")),
        (label(lang, "Agri yield status"), hazard_summary_value(main_row, child_df, "agri_yield_status")),
        (label(lang, "Agri yield reasons"), hazard_summary_value(main_row, child_df, "agri_yield_reasons")),
        (label(lang, "Agriculture vulnerability factors"), hazard_summary_value(main_row, child_df, "agri_vulnerability_factors")),
        (label(lang, "Most vulnerable groups in agriculture"), hazard_summary_value(main_row, child_df, "agri_most_vulnerable_groups")),
        (label(lang, "Livestock yield status"), hazard_summary_value(main_row, child_df, "livestock_yield_status")),
        (label(lang, "Livestock yield reasons"), hazard_summary_value(main_row, child_df, "livestock_yield_reasons")),
        (label(lang, "Livestock vulnerability factors"), hazard_summary_value(main_row, child_df, "livestock_vulnerability_factors")),
        (label(lang, "Vulnerable household count"), hazard_summary_value(main_row, child_df, "vulnerable_household_count")),
        (label(lang, "Vulnerable root causes"), hazard_summary_value(main_row, child_df, "vulnerable_root_causes")),
        (label(lang, "Climate change pattern"), hazard_summary_value(main_row, child_df, "climate_change_pattern")),
        (label(lang, "Pest / disease pattern"), hazard_summary_value(main_row, child_df, "pest_disease_pattern")),
    ]
    if any(v != "-" for _, v in summary_pairs):
        add_heading(doc, tr(lang, "hazard_summary"), level=2)
        add_kv_table(doc, summary_pairs)

    if child_df.empty:
        doc.add_paragraph(tr(lang, "no_data"))
        return

    for i, (_, row) in enumerate(child_df.iterrows(), start=1):
        add_heading(doc, tr(lang, "hazard").format(i=i, name=clean_text(getv(row, 'hazard_type'))), level=2)
        add_kv_table(doc, [
            (label(lang, "Hazard type"), getv(row, "hazard_type")),
            (label(lang, "Date / period"), getv(row, "hazard_date_or_period")),
        ])
        for display_key, canonical in [
            ("Hazard description", "hazard_description"), ("Physical impact", "physical_impact"), ("Psychosocial impact", "psychosocial_impact"),
            ("Economic impact", "economic_impact"), ("Environmental impact", "environmental_impact"), ("Land / water impact", "land_water_impact"),
            ("Household expenditure impact", "household_expenditure_impact"), ("Agriculture impact", "agriculture_impact"),
            ("Livestock impact", "livestock_impact"), ("Vulnerable group impact", "vulnerable_group_impact"),
            ("Community group impact", "community_group_impact"), ("Community recovery actions", "community_recovery_actions")
        ]:
            add_paragraph_text(doc, label(lang, display_key), getv(row, canonical))
        doc.add_paragraph()

    village_level_extras = [
        ("Coping mechanisms", "coping_mechanisms"),
    ]
    for display_key, canonical in village_level_extras:
        value = hazard_summary_value(main_row, child_df, canonical)
        if value != "-":
            add_paragraph_text(doc, label(lang, display_key), value)
            doc.add_paragraph()


def section_seasonal_calendar(doc, child_df, lang: str):
    add_heading(doc, tr(lang, "sec6"), level=1)
    if child_df.empty:
        doc.add_paragraph(tr(lang, "no_data"))
        return
    records = []
    for _, row in child_df.iterrows():
        records.append({
            tr(lang, "category"): clean_text(getv(row, "calendar_category")),
            tr(lang, "item_name"): clean_text(getv(row, "calendar_item_name")),
            tr(lang, "active_months"): month_list_from_row(row, "active_months"),
            tr(lang, "peak_months"): month_list_from_row(row, "peak_months"),
            tr(lang, "low_months"): month_list_from_row(row, "low_months"),
            tr(lang, "notes"): clean_text(getv(row, "calendar_notes")),
        })
    add_df_table(doc, pd.DataFrame(records), empty_text=tr(lang, "no_data"))


def section_priority_ranking(doc, child_df, lang: str):
    add_heading(doc, tr(lang, "sec7"), level=1)
    if child_df.empty:
        doc.add_paragraph(tr(lang, "no_data"))
        return
    df = child_df.copy()
    sort_col = getcol(df, "rank_order")
    score_col = getcol(df, "score")
    if sort_col:
        if score_col:
            df = df.sort_values(by=[sort_col, score_col], ascending=[True, False])
        else:
            df = df.sort_values(by=[sort_col], ascending=[True])
    records = []
    for _, row in df.iterrows():
        records.append({
            tr(lang, "category"): clean_text(getv(row, "ranking_category")),
            tr(lang, "option"): clean_text(getv(row, "option_name")),
            tr(lang, "score"): clean_text(getv(row, "score")),
            tr(lang, "rank"): clean_text(getv(row, "rank_order")),
            tr(lang, "top3"): clean_text(getv(row, "selected_top3")),
            tr(lang, "needs_relevance"): clean_text(getv(row, "needs_relevance")),
            tr(lang, "feasibility"): clean_text(getv(row, "feasibility")),
            tr(lang, "practicality"): clean_text(getv(row, "practicality")),
            tr(lang, "vulnerable_inclusion"): clean_text(getv(row, "vulnerable_inclusion")),
            tr(lang, "short_term_feasible"): clean_text(getv(row, "short_term_feasible")),
            tr(lang, "social_protection_link"): clean_text(getv(row, "social_protection_link")),
            tr(lang, "community_feedback"): clean_text(getv(row, "community_feedback")),
        })
    add_df_table(doc, pd.DataFrame(records), empty_text=tr(lang, "no_data"))


def build_village_doc(workbook: Dict[str, pd.DataFrame], row: pd.Series, lang: str) -> bytes:
    doc = base_document()
    add_heading(doc, tr(lang, "report_title"), level=0)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run(tr(lang, "subtitle").format(village=village_name(row), township=clean_text(getv(row, 'township_name'))))
    r.font.name = "Noto Sans Myanmar"
    r.font.size = Pt(11)
    r.italic = True
    doc.add_paragraph()

    parent_index = row.get("_index")
    children = {}
    for sheet_name in REPEAT_SHEETS:
        df = workbook.get(sheet_name, pd.DataFrame())
        if "_parent_index" in df.columns:
            children[sheet_name] = df[df["_parent_index"] == parent_index].copy()
        else:
            children[sheet_name] = pd.DataFrame()

    section_village_general(doc, row, lang)
    section_livelihood(doc, row, children["livelihood_repeat"], lang)
    section_agri(doc, row, children["agriculture_repeat"], lang)
    section_livestock(doc, row, children["livestock_repeat"], lang)
    section_hazard(doc, row, children["hazard_repeat"], lang)
    section_seasonal_calendar(doc, children["seasonal_calendar_repeat"], lang)
    section_priority_ranking(doc, children["priority_ranking_repeat"], lang)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def generate_reports(workbook: Dict[str, pd.DataFrame]):
    main_sheet = find_main_sheet(workbook)
    if not main_sheet:
        raise ValueError(f"Main sheet not found. Expected one of: {', '.join(MAIN_SHEET_CANDIDATES)}")
    main_df = workbook[main_sheet].copy()
    if "_index" not in main_df.columns:
        raise ValueError("Main sheet must contain '_index' column.")
    lang = detect_language(main_df)
    main_df = main_df.sort_values(by="_index")
    reports = []
    for _, row in main_df.iterrows():
        reports.append({
            "village_name": village_name(row),
            "parent_index": row.get("_index"),
            "bytes": build_village_doc(workbook, row, lang),
            "language": lang,
        })
    return reports


def build_zip(reports):
    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        for item in reports:
            zf.writestr(sanitize_filename(item["village_name"]) + ".docx", item["bytes"])
    buffer.seek(0)
    return buffer.getvalue()
